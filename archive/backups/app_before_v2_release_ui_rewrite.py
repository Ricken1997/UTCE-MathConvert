#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from utce_core_refactor_work import (
    analyze_lines,
    build_latex_output,
    build_diagnosis_from_severity_counts,
)

from omml_converter import convert_lines_to_omml


APP_TITLE = "UTCE MathConvert v2.0 Product Beta Free"
WINDOW_SIZE = "1450x1000"
TEXT_HEIGHT = 8
EXPAND_SIZE = "1300x850"


# ============================================================
# Word / OMML Writer
# ============================================================

def normalize_omml_xml(xml: str) -> str:
    xml = xml.strip()

    if not xml:
        return ""

    if "xmlns:m=" not in xml and xml.startswith("<m:oMathPara"):
        xml = xml.replace(
            "<m:oMathPara",
            f"<m:oMathPara {nsdecls('m', 'w')}",
            1
        )

    return xml


def insert_omml_paragraph(doc: Document, omml_xml: str) -> bool:
    try:
        normalized = normalize_omml_xml(omml_xml)
        element = parse_xml(normalized)

        paragraph = doc.add_paragraph()
        paragraph._p.append(element)
        return True

    except Exception as e:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"[OMML INSERT ERROR] {e}")
        paragraph.add_run("\n")
        paragraph.add_run(omml_xml)
        return False


def save_word():
    plain_text = input_box.get("1.0", tk.END).strip()

    if not plain_text:
        messagebox.showwarning("No Input", "Please enter text before saving Word file.")
        return

    path = filedialog.asksaveasfilename(
        title="Save Word File",
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )

    if not path:
        return

    try:
        input_lines = plain_text.splitlines()
        omml_lines = convert_lines_to_omml(input_lines)

        doc = Document()

        for omml_xml in omml_lines:
            if omml_xml.strip():
                insert_omml_paragraph(doc, omml_xml)

        doc.save(path)
        messagebox.showinfo("Saved", "Word document saved successfully.")

    except Exception as e:
        messagebox.showerror("Save Word Error", str(e))


# ============================================================
# Convert
# ============================================================

def convert_text():
    plain_text = input_box.get("1.0", tk.END).strip()

    if not plain_text:
        messagebox.showwarning("No Input", "Please enter text before converting.")
        return

    input_lines = plain_text.splitlines()

    latex_lines, warnings, severity_counts = analyze_lines(input_lines)
    diagnosis = build_diagnosis_from_severity_counts(severity_counts)

    latex_output = build_latex_output(latex_lines, output_mode_var.get())
    omml_lines = convert_lines_to_omml(input_lines)

    output_box.delete("1.0", tk.END)
    output_box.insert(tk.END, latex_output)

    mathml_box.delete("1.0", tk.END)
    mathml_box.insert(tk.END, "<math><mtext>OMML generated</mtext></math>")

    omml_box.delete("1.0", tk.END)
    omml_box.insert(tk.END, "\n".join(omml_lines))

    warning_box.config(text="\n".join(warnings) if warnings else "No warnings.")

    confidence_var.set(f"Confidence: {diagnosis.confidence_score:.1f}")
    risk_var.set(f"Predictive Risk: {diagnosis.predictive_risk:.1f}")
    level_var.set(f"Risk Level: {diagnosis.risk_level}")


# ============================================================
# Large Editor / Viewer
# ============================================================

def open_large_editor(title: str, text_widget: tk.Text, editable: bool = True):
    window = tk.Toplevel(root)
    window.title(title)
    window.geometry(EXPAND_SIZE)

    frame = tk.Frame(window, padx=8, pady=8)
    frame.pack(fill=tk.BOTH, expand=True)

    label = tk.Label(frame, text=title, font=("Helvetica", 16, "bold"), anchor="w")
    label.pack(fill=tk.X)

    text_frame = tk.Frame(frame)
    text_frame.pack(fill=tk.BOTH, expand=True, pady=8)

    y_scroll = tk.Scrollbar(text_frame, orient=tk.VERTICAL)
    x_scroll = tk.Scrollbar(text_frame, orient=tk.HORIZONTAL)

    large_text = tk.Text(
        text_frame,
        wrap="none",
        undo=True,
        yscrollcommand=y_scroll.set,
        xscrollcommand=x_scroll.set,
    )

    y_scroll.config(command=large_text.yview)
    x_scroll.config(command=large_text.xview)

    large_text.grid(row=0, column=0, sticky="nsew")
    y_scroll.grid(row=0, column=1, sticky="ns")
    x_scroll.grid(row=1, column=0, sticky="ew")

    text_frame.rowconfigure(0, weight=1)
    text_frame.columnconfigure(0, weight=1)

    large_text.insert("1.0", text_widget.get("1.0", tk.END))

    if not editable:
        large_text.config(state=tk.DISABLED)

    button_frame = tk.Frame(frame)
    button_frame.pack(fill=tk.X)

    if editable:
        def apply_back():
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", large_text.get("1.0", tk.END))
            window.destroy()

        tk.Button(button_frame, text="Apply Back", command=apply_back).pack(side=tk.LEFT, padx=4)

    tk.Button(button_frame, text="Close", command=window.destroy).pack(side=tk.LEFT, padx=4)


# ============================================================
# Reusable Text Area
# ============================================================

def create_text_area(parent, title: str, height: int = TEXT_HEIGHT, editable: bool = True):
    section = tk.Frame(parent)
    section.pack(fill=tk.BOTH, expand=False, pady=(8, 0))

    header = tk.Frame(section)
    header.pack(fill=tk.X)

    tk.Label(header, text=title, anchor="w").pack(side=tk.LEFT)

    text_frame = tk.Frame(section)
    text_frame.pack(fill=tk.BOTH, expand=True)

    y_scroll = tk.Scrollbar(text_frame, orient=tk.VERTICAL)
    x_scroll = tk.Scrollbar(text_frame, orient=tk.HORIZONTAL)

    text_widget = tk.Text(
        text_frame,
        height=height,
        wrap="none",
        undo=True,
        yscrollcommand=y_scroll.set,
        xscrollcommand=x_scroll.set,
    )

    y_scroll.config(command=text_widget.yview)
    x_scroll.config(command=text_widget.xview)

    text_widget.grid(row=0, column=0, sticky="nsew")
    y_scroll.grid(row=0, column=1, sticky="ns")
    x_scroll.grid(row=1, column=0, sticky="ew")

    text_frame.rowconfigure(0, weight=1)
    text_frame.columnconfigure(0, weight=1)

    if not editable:
        text_widget.config(state=tk.NORMAL)

    tk.Button(
        header,
        text="Expand",
        command=lambda: open_large_editor(title, text_widget, editable=editable)
    ).pack(side=tk.RIGHT)

    return text_widget


# ============================================================
# File / UI Actions
# ============================================================

def open_text_file():
    path = filedialog.askopenfilename(
        title="Open Text File",
        filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
    )

    if not path:
        return

    text = Path(path).read_text(encoding="utf-8")
    input_box.delete("1.0", tk.END)
    input_box.insert(tk.END, text)


def example_input():
    sample = """align(
frac(sum(i,1,n,pow(a,2)),sqrt(prod(j,1,m,b))),int(x,0,1,frac(alpha,beta));
cases(frac(a,b),x>0;sqrt(c),x=0;sum(i,1,n,a),x<0),lim(n,inf,frac(1,n));
matrix(alpha,beta,gamma;delta,epsilon,zeta;eta,theta,lambda),frac(int(t,0,pi,sin(t)),sum(k,1,n,pow(k,2)))
)"""
    input_box.delete("1.0", tk.END)
    input_box.insert(tk.END, sample)


def save_output():
    path = filedialog.asksaveasfilename(
        title="Save Output",
        defaultextension=".txt",
        filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
    )

    if not path:
        return

    content = output_box.get("1.0", tk.END).strip()
    Path(path).write_text(content, encoding="utf-8")
    messagebox.showinfo("Saved", "Output saved successfully.")


def copy_output():
    text = output_box.get("1.0", tk.END).strip()
    root.clipboard_clear()
    root.clipboard_append(text)
    root.update()
    messagebox.showinfo("Copied", "Output copied to clipboard.")


def clear_all():
    input_box.delete("1.0", tk.END)
    output_box.delete("1.0", tk.END)
    mathml_box.delete("1.0", tk.END)
    omml_box.delete("1.0", tk.END)

    warning_box.config(text="")
    confidence_var.set("Confidence:")
    risk_var.set("Predictive Risk:")
    level_var.set("Risk Level:")


# ============================================================
# GUI
# ============================================================

root = tk.Tk()
root.title(APP_TITLE)
root.geometry(WINDOW_SIZE)

main_frame = tk.Frame(root, padx=12, pady=12)
main_frame.pack(fill=tk.BOTH, expand=True)

title_label = tk.Label(
    main_frame,
    text=APP_TITLE,
    font=("Helvetica", 22, "bold"),
    anchor="w"
)
title_label.pack(fill=tk.X)

file_type_label = tk.Label(
    main_frame,
    text="Detected File Type: None",
    anchor="w"
)
file_type_label.pack(fill=tk.X, pady=(8, 0))

input_box = create_text_area(main_frame, "Input", height=TEXT_HEIGHT, editable=True)

button_frame = tk.Frame(main_frame)
button_frame.pack(fill=tk.X, pady=10)

tk.Button(button_frame, text="Convert", command=convert_text).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Open Text File", command=open_text_file).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Example Input", command=example_input).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Save Output", command=save_output).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Save Word", command=save_word).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Copy Output", command=copy_output).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Clear", command=clear_all).pack(side=tk.LEFT, padx=4)

output_mode_var = tk.StringVar(value="inline")

tk.Label(button_frame, text="Output Mode:").pack(side=tk.LEFT, padx=(30, 4))
tk.Radiobutton(button_frame, text="Inline", variable=output_mode_var, value="inline").pack(side=tk.LEFT)
tk.Radiobutton(button_frame, text="Block", variable=output_mode_var, value="block").pack(side=tk.LEFT)

output_box = create_text_area(main_frame, "LaTeX Output", height=TEXT_HEIGHT, editable=True)
mathml_box = create_text_area(main_frame, "MathML Output", height=TEXT_HEIGHT, editable=True)
omml_box = create_text_area(main_frame, "OMML Output", height=TEXT_HEIGHT, editable=True)

tk.Label(main_frame, text="Warnings", anchor="w").pack(fill=tk.X, pady=(8, 0))
warning_box = tk.Label(main_frame, text="", anchor="w", justify="left")
warning_box.pack(fill=tk.X)

tk.Label(main_frame, text="Structural Diagnosis", anchor="w").pack(fill=tk.X, pady=(12, 0))

confidence_var = tk.StringVar(value="Confidence:")
risk_var = tk.StringVar(value="Predictive Risk:")
level_var = tk.StringVar(value="Risk Level:")

tk.Label(main_frame, textvariable=confidence_var, anchor="w").pack(fill=tk.X, padx=40)
tk.Label(main_frame, textvariable=risk_var, anchor="w").pack(fill=tk.X, padx=40)
tk.Label(main_frame, textvariable=level_var, anchor="w").pack(fill=tk.X, padx=40)

root.mainloop()