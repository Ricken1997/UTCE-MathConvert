#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path
import time
import tempfile
import subprocess

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from utce_core_refactor_work import (
    analyze_lines,
    build_latex_output,
    build_diagnosis_from_severity_counts,
)

from omml_converter import convert_lines_to_omml


APP_TITLE = "UTCE MathConvert v2.0 Product Release Free"
WINDOW_SIZE = "1450x1000"
EXPAND_SIZE = "1300x850"
PREVIEW_SIZE = "1200x800"
TEXT_HEIGHT = 8


# ============================================================
# OMML / Word
# ============================================================

def normalize_omml_xml(xml: str) -> str:
    xml = xml.strip()
    if not xml:
        return ""

    if "xmlns:m=" not in xml and xml.startswith("<m:oMathPara"):
        xml = xml.replace(
            "<m:oMathPara",
            f"<m:oMathPara {nsdecls('m', 'w')}",
            1
        )
    return xml


def insert_omml_paragraph(doc: Document, omml_xml: str) -> bool:
    try:
        normalized = normalize_omml_xml(omml_xml)
        element = parse_xml(normalized)
        paragraph = doc.add_paragraph()
        paragraph._p.append(element)
        return True
    except Exception as e:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"[OMML INSERT ERROR] {e}")
        paragraph.add_run("\n")
        paragraph.add_run(omml_xml)
        return False


def build_word_document_from_input() -> Document:
    plain_text = input_box.get("1.0", tk.END).strip()
    input_lines = plain_text.splitlines()
    omml_lines = convert_lines_to_omml(input_lines)

    doc = Document()

    for omml_xml in omml_lines:
        if omml_xml.strip():
            insert_omml_paragraph(doc, omml_xml)

    return doc


def get_current_summary() -> str:
    return (
        f"Expressions: {expr_var.get().replace('Expressions:', '').strip() or 'Unknown'}\n"
        f"Warnings: {current_warning_count.get()}\n"
        f"{confidence_var.get()}\n"
        f"{risk_var.get()}\n"
        f"{level_var.get()}\n"
        f"{time_var.get()}"
    )


# ============================================================
# Preview Dialogs
# ============================================================

def create_scrolled_text(parent, height=20, editable=True):
    frame = tk.Frame(parent)

    y_scroll = tk.Scrollbar(frame, orient=tk.VERTICAL)
    x_scroll = tk.Scrollbar(frame, orient=tk.HORIZONTAL)

    text = tk.Text(
        frame,
        height=height,
        wrap="none",
        undo=True,
        yscrollcommand=y_scroll.set,
        xscrollcommand=x_scroll.set,
    )

    y_scroll.config(command=text.yview)
    x_scroll.config(command=text.xview)

    text.grid(row=0, column=0, sticky="nsew")
    y_scroll.grid(row=0, column=1, sticky="ns")
    x_scroll.grid(row=1, column=0, sticky="ew")

    frame.rowconfigure(0, weight=1)
    frame.columnconfigure(0, weight=1)

    if not editable:
        text.config(state=tk.DISABLED)

    return frame, text


def preview_text_window(title: str, content: str, on_confirm=None, confirm_label="OK"):
    window = tk.Toplevel(root)
    window.title(title)
    window.geometry(PREVIEW_SIZE)

    frame = tk.Frame(window, padx=10, pady=10)
    frame.pack(fill=tk.BOTH, expand=True)

    tk.Label(frame, text=title, font=("Helvetica", 16, "bold"), anchor="w").pack(fill=tk.X)

    text_frame, text_widget = create_scrolled_text(frame, height=28, editable=True)
    text_frame.pack(fill=tk.BOTH, expand=True, pady=8)

    text_widget.insert("1.0", content)

    button_frame = tk.Frame(frame)
    button_frame.pack(fill=tk.X)

    def confirm():
        if on_confirm:
            on_confirm(text_widget.get("1.0", tk.END).strip())
        window.destroy()

    if on_confirm:
        tk.Button(button_frame, text=confirm_label, command=confirm).pack(side=tk.LEFT, padx=4)

    tk.Button(button_frame, text="Close", command=window.destroy).pack(side=tk.LEFT, padx=4)


def output_selection_dialog(action: str):
    window = tk.Toplevel(root)
    window.title(action)
    window.geometry("420x260")

    frame = tk.Frame(window, padx=14, pady=14)
    frame.pack(fill=tk.BOTH, expand=True)

    tk.Label(frame, text=f"{action}: select output type", font=("Helvetica", 14, "bold")).pack(anchor="w")

    selected = tk.StringVar(value="LaTeX")

    options = [
        ("LaTeX", output_box),
        ("MathML", mathml_box),
        ("OMML", omml_box),
    ]

    for label, _widget in options:
        tk.Radiobutton(frame, text=label, variable=selected, value=label).pack(anchor="w", pady=4)

    def proceed():
        kind = selected.get()
        widget = dict(options)[kind]
        content = widget.get("1.0", tk.END).strip()

        if not content:
            messagebox.showwarning("No Output", f"No {kind} output available.")
            return

        window.destroy()

        if action == "Save Output":
            def save_confirm(edited_content: str):
                path = filedialog.asksaveasfilename(
                    title=f"Save {kind}",
                    defaultextension=".txt",
                    filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
                )
                if not path:
                    return
                Path(path).write_text(edited_content, encoding="utf-8")
                messagebox.showinfo("Saved", f"{kind} output saved successfully.")

            preview_text_window(
                f"{kind} Preview Before Save",
                content,
                on_confirm=save_confirm,
                confirm_label="Save"
            )

        elif action == "Copy Output":
            def copy_confirm(edited_content: str):
                root.clipboard_clear()
                root.clipboard_append(edited_content)
                root.update()
                messagebox.showinfo("Copied", f"{kind} output copied to clipboard.")

            preview_text_window(
                f"{kind} Preview Before Copy",
                content,
                on_confirm=copy_confirm,
                confirm_label="Copy"
            )

    tk.Button(frame, text="Preview", command=proceed).pack(side=tk.LEFT, padx=4, pady=18)
    tk.Button(frame, text="Cancel", command=window.destroy).pack(side=tk.LEFT, padx=4, pady=18)


def open_file(path: str):
    try:
        subprocess.Popen(["open", path])
    except Exception as e:
        messagebox.showerror("Open Error", str(e))


def open_folder(path: str):
    try:
        folder = str(Path(path).parent)
        subprocess.Popen(["open", folder])
    except Exception as e:
        messagebox.showerror("Open Folder Error", str(e))


def saved_word_dialog(path: str):
    window = tk.Toplevel(root)
    window.title("Word Saved")
    window.geometry("520x260")

    frame = tk.Frame(window, padx=14, pady=14)
    frame.pack(fill=tk.BOTH, expand=True)

    tk.Label(
        frame,
        text="Word document saved successfully.",
        font=("Helvetica", 15, "bold"),
        anchor="w"
    ).pack(fill=tk.X)

    tk.Label(
        frame,
        text=f"Saved to:\n{path}",
        justify="left",
        anchor="w"
    ).pack(fill=tk.X, pady=12)

    button_frame = tk.Frame(frame)
    button_frame.pack(fill=tk.X, pady=10)

    tk.Button(
        button_frame,
        text="Open Word",
        command=lambda: open_file(path)
    ).pack(side=tk.LEFT, padx=4)

    tk.Button(
        button_frame,
        text="Open Folder",
        command=lambda: open_folder(path)
    ).pack(side=tk.LEFT, padx=4)

    tk.Button(
        button_frame,
        text="Close",
        command=window.destroy
    ).pack(side=tk.LEFT, padx=4)


def save_word_preview_dialog():
    plain_text = input_box.get("1.0", tk.END).strip()

    if not plain_text:
        messagebox.showwarning("No Input", "Please enter text before saving Word file.")
        return

    window = tk.Toplevel(root)
    window.title("Save Word Preview")
    window.geometry("560x420")

    frame = tk.Frame(window, padx=14, pady=14)
    frame.pack(fill=tk.BOTH, expand=True)

    tk.Label(frame, text="Word Output Preview", font=("Helvetica", 16, "bold")).pack(anchor="w")

    summary = (
        "A Word document will be generated from the current input.\n\n"
        + get_current_summary()
        + "\n\nChoose an action:"
    )

    tk.Label(frame, text=summary, justify="left", anchor="w").pack(fill=tk.X, pady=12)

    def preview_word():
        try:
            doc = build_word_document_from_input()
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tmp_path = tmp.name
            tmp.close()

            doc.save(tmp_path)
            subprocess.Popen(["open", tmp_path])
        except Exception as e:
            messagebox.showerror("Preview Error", str(e))

    def save_word_final():
        path = filedialog.asksaveasfilename(
            title="Save Word File",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )

        if not path:
            return

        try:
            doc = build_word_document_from_input()
            doc.save(path)
            saved_word_dialog(path)
            window.destroy()
        except Exception as e:
            messagebox.showerror("Save Word Error", str(e))

    button_frame = tk.Frame(frame)
    button_frame.pack(fill=tk.X, pady=10)

    tk.Button(button_frame, text="Preview Word", command=preview_word).pack(side=tk.LEFT, padx=4)
    tk.Button(button_frame, text="Save", command=save_word_final).pack(side=tk.LEFT, padx=4)
    tk.Button(button_frame, text="Cancel", command=window.destroy).pack(side=tk.LEFT, padx=4)


# ============================================================
# Convert
# ============================================================

def convert_text():
    plain_text = input_box.get("1.0", tk.END).strip()

    if not plain_text:
        messagebox.showwarning("No Input", "Please enter text before converting.")
        return

    status_var.set("Converting...")
    root.update_idletasks()

    start_time = time.time()
    input_lines = plain_text.splitlines()

    try:
        latex_lines, warnings, severity_counts = analyze_lines(input_lines)
        diagnosis = build_diagnosis_from_severity_counts(severity_counts)

        latex_output = build_latex_output(latex_lines, output_mode_var.get())
        omml_lines = convert_lines_to_omml(input_lines)

        elapsed_ms = int((time.time() - start_time) * 1000)

        output_box.delete("1.0", tk.END)
        output_box.insert(tk.END, latex_output)

        mathml_box.delete("1.0", tk.END)
        mathml_box.insert(tk.END, "<math><mtext>OMML generated</mtext></math>")

        omml_box.delete("1.0", tk.END)
        omml_box.insert(tk.END, "\n".join(omml_lines))

        warning_box.config(text="\n".join(warnings) if warnings else "No warnings.")

        current_warning_count.set(len(warnings))

        confidence_var.set(f"Confidence: {diagnosis.confidence_score:.1f}")
        risk_var.set(f"Predictive Risk: {diagnosis.predictive_risk:.1f}")
        level_var.set(f"Risk Level: {diagnosis.risk_level}")
        time_var.set(f"Parser Time: {elapsed_ms} ms")
        expr_var.set(f"Expressions: {len(omml_lines)}")

        status_var.set("Ready")

    except Exception as e:
        status_var.set("Error")
        messagebox.showerror("Convert Error", str(e))


# ============================================================
# Large Editor
# ============================================================

def open_large_editor(title: str, text_widget: tk.Text, editable: bool = True):
    window = tk.Toplevel(root)
    window.title(title)
    window.geometry(EXPAND_SIZE)

    frame = tk.Frame(window, padx=8, pady=8)
    frame.pack(fill=tk.BOTH, expand=True)

    tk.Label(frame, text=title, font=("Helvetica", 16, "bold"), anchor="w").pack(fill=tk.X)

    text_frame, large_text = create_scrolled_text(frame, height=30, editable=True)
    text_frame.pack(fill=tk.BOTH, expand=True, pady=8)

    large_text.insert("1.0", text_widget.get("1.0", tk.END))

    if not editable:
        large_text.config(state=tk.DISABLED)

    button_frame = tk.Frame(frame)
    button_frame.pack(fill=tk.X)

    if editable:
        def apply_back():
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", large_text.get("1.0", tk.END))
            window.destroy()

        tk.Button(button_frame, text="Apply Back", command=apply_back).pack(side=tk.LEFT, padx=4)

    tk.Button(button_frame, text="Close", command=window.destroy).pack(side=tk.LEFT, padx=4)


def create_text_area(parent, title: str, height: int = TEXT_HEIGHT, editable: bool = True):
    section = tk.Frame(parent)
    section.pack(fill=tk.BOTH, expand=False, pady=(8, 0))

    header = tk.Frame(section)
    header.pack(fill=tk.X)

    tk.Label(header, text=title, anchor="w").pack(side=tk.LEFT)

    text_frame, text_widget = create_scrolled_text(section, height=height, editable=True)
    text_frame.pack(fill=tk.BOTH, expand=True)

    tk.Button(
        header,
        text="Expand",
        command=lambda: open_large_editor(title, text_widget, editable=editable)
    ).pack(side=tk.RIGHT)

    return text_widget


# ============================================================
# File / UI Actions
# ============================================================

def open_text_file():
    path = filedialog.askopenfilename(
        title="Open Text File",
        filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
    )

    if not path:
        return

    text = Path(path).read_text(encoding="utf-8")
    input_box.delete("1.0", tk.END)
    input_box.insert(tk.END, text)


def example_input():
    sample = """align(
frac(sum(i,1,n,pow(a,2)),sqrt(prod(j,1,m,b))),int(x,0,1,frac(alpha,beta));
cases(frac(a,b),x>0;sqrt(c),x=0;sum(i,1,n,a),x<0),lim(n,inf,frac(1,n));
matrix(alpha,beta,gamma;delta,epsilon,zeta;eta,theta,lambda),frac(int(t,0,pi,sin(t)),sum(k,1,n,pow(k,2)))
)"""
    input_box.delete("1.0", tk.END)
    input_box.insert(tk.END, sample)


def save_output():
    output_selection_dialog("Save Output")


def copy_output():
    output_selection_dialog("Copy Output")


def clear_all():
    if not messagebox.askyesno("Confirm Clear", "Clear all input and output fields?"):
        return

    input_box.delete("1.0", tk.END)
    output_box.delete("1.0", tk.END)
    mathml_box.delete("1.0", tk.END)
    omml_box.delete("1.0", tk.END)

    warning_box.config(text="")
    confidence_var.set("Confidence:")
    risk_var.set("Predictive Risk:")
    level_var.set("Risk Level:")
    time_var.set("Parser Time:")
    expr_var.set("Expressions:")
    current_warning_count.set(0)
    status_var.set("Ready")


# ============================================================
# GUI
# ============================================================

root = tk.Tk()
root.title(APP_TITLE)
root.geometry(WINDOW_SIZE)

current_warning_count = tk.IntVar(value=0)

main_frame = tk.Frame(root, padx=12, pady=12)
main_frame.pack(fill=tk.BOTH, expand=True)

tk.Label(
    main_frame,
    text=APP_TITLE,
    font=("Helvetica", 22, "bold"),
    anchor="w"
).pack(fill=tk.X)

status_var = tk.StringVar(value="Ready")
tk.Label(main_frame, textvariable=status_var, anchor="w").pack(fill=tk.X, pady=(4, 0))

tk.Label(main_frame, text="Detected File Type: None", anchor="w").pack(fill=tk.X, pady=(4, 0))

input_box = create_text_area(main_frame, "Input", height=TEXT_HEIGHT, editable=True)

button_frame = tk.Frame(main_frame)
button_frame.pack(fill=tk.X, pady=10)

tk.Button(button_frame, text="Convert", command=convert_text).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Open Text File", command=open_text_file).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Example Input", command=example_input).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Save Output", command=save_output).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Save Word", command=save_word_preview_dialog).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Copy Output", command=copy_output).pack(side=tk.LEFT, padx=4)
tk.Button(button_frame, text="Clear", command=clear_all).pack(side=tk.LEFT, padx=4)

output_mode_var = tk.StringVar(value="inline")

tk.Label(button_frame, text="Output Mode:").pack(side=tk.LEFT, padx=(30, 4))
tk.Radiobutton(button_frame, text="Inline", variable=output_mode_var, value="inline").pack(side=tk.LEFT)
tk.Radiobutton(button_frame, text="Block", variable=output_mode_var, value="block").pack(side=tk.LEFT)

output_box = create_text_area(main_frame, "LaTeX Output", height=TEXT_HEIGHT, editable=True)
mathml_box = create_text_area(main_frame, "MathML Output", height=TEXT_HEIGHT, editable=True)
omml_box = create_text_area(main_frame, "OMML Output", height=TEXT_HEIGHT, editable=True)

tk.Label(main_frame, text="Warnings", anchor="w").pack(fill=tk.X, pady=(8, 0))
warning_box = tk.Label(main_frame, text="", anchor="w", justify="left")
warning_box.pack(fill=tk.X)

tk.Label(main_frame, text="Structural Diagnosis", anchor="w").pack(fill=tk.X, pady=(12, 0))

confidence_var = tk.StringVar(value="Confidence:")
risk_var = tk.StringVar(value="Predictive Risk:")
level_var = tk.StringVar(value="Risk Level:")
time_var = tk.StringVar(value="Parser Time:")
expr_var = tk.StringVar(value="Expressions:")

tk.Label(main_frame, textvariable=confidence_var, anchor="w").pack(fill=tk.X, padx=40)
tk.Label(main_frame, textvariable=risk_var, anchor="w").pack(fill=tk.X, padx=40)
tk.Label(main_frame, textvariable=level_var, anchor="w").pack(fill=tk.X, padx=40)
tk.Label(main_frame, textvariable=time_var, anchor="w").pack(fill=tk.X, padx=40)
tk.Label(main_frame, textvariable=expr_var, anchor="w").pack(fill=tk.X, padx=40)

root.mainloop()